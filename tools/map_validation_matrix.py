"""Classify the validation matrix against the 168-page PDF-faithful baseline."""

from __future__ import annotations

import json
import re
from pathlib import Path

from docx import Document


ROOT = Path(__file__).resolve().parents[1]
REPORT = Path(r"C:\Users\Admin\Downloads\RIPOTI YA ADT VALIDATION SAYANSI 11.08.2026.docx")


def printed_pages(text: str, allow_trailing: bool = False) -> list[int]:
    found: set[int] = set()
    pattern = r"\b(?:uk(?:urasa)?|ukurasa)\.?\s*(\d{1,3}(?:\s*(?:,|/|&|na|[-–])\s*\d{1,3})*)"
    for match in re.finditer(pattern, text, flags=re.IGNORECASE):
        values = [int(value) for value in re.findall(r"\d{1,3}", match.group(1))]
        if "-" in match.group(1) or "–" in match.group(1):
            if len(values) == 2 and values[1] >= values[0] and values[1] - values[0] <= 30:
                found.update(range(values[0], values[1] + 1))
            else:
                found.update(values)
        else:
            found.update(values)
    if allow_trailing and not found:
        trailing = re.search(r"(?:^|\s)(\d{1,3})\s*[.)]?\s*$", text.strip())
        if trailing and 1 <= int(trailing.group(1)) <= 162:
            found.add(int(trailing.group(1)))
    return sorted(page for page in found if 1 <= page <= 162)


def classify(cells: list[str]) -> tuple[str, str, str]:
    joined = " ".join(cells).lower()
    if "chemsha bongo" in joined or "chemshabongo" in joined or re.search(r"\bqz\d+\.html", joined):
        return "resolved_obsolete", "quiz", "Chemsha bongo/qz already removed with the old extra pages."
    if any(term in joined for term in ["kujibia", "sehemu ya kujibu", "sehemu ya kuweka majibu", "uingizaji wa majibu", "kupokea majibu"]):
        return "hold", "answer_input", "Answer-space/input changes are explicitly out of scope for now."
    if any(term in joined for term in ["option ya tuma", "option ya kutuma", "chaguo la tuma", "submit"]):
        return "resolved_obsolete", "old_ui_control", "The PDF-faithful baseline no longer contains the old submit control."
    if "kurasa zote" in joined:
        return "global", "book_wide", "Apply through the shared accessible-text/runtime layer."
    if "for online reading only" in joined or "head footer" in joined or "foot note" in joined or "fotot note" in joined:
        return "global", "speech_cleanup", "Remove from speech/transcript without changing the PDF image."
    if any(term in joined for term in ["quorum", "program saidizi", "programu saidizi", "kipima joto sauti", "saa ya mtetemo"]):
        return "page_specific", "assistive_technology", "Add an accessible alternative after verifying the current activity and page."
    if any(term in joined for term in ["lugha jumuishi", "vitenzi jumuishi", "kuchunguza", "kuhisi", "kugusa", "kubaini"]):
        return "page_specific", "inclusive_language", "Apply to accessible text; retain the photographed PDF wording."
    if any(term in joined for term in ["maelezo ya picha", "maelezo ya kielelezo", "kielelezo kina"]):
        return "page_specific", "figure_description", "Add a concise screen-reader description tied to the matching figure."
    if any(term in joined for term in ["matamshi", "itamkwe", "sauti", "audio", "kisoma skrini"]):
        return "page_specific", "speech_pronunciation", "Correct the accessible transcript/speech form after checking the source wording."
    if any(term in joined for term in ["imeongezwa", "hayapo katika kitabu original", "haipo katika original"]):
        return "resolved_obsolete", "old_extra_content", "The PDF-faithful rebuild removed content that was not in the source PDF."
    return "page_specific", "content_or_layout", "Verify against the current PDF page before applying."


def reference_pages(reference: str) -> list[int]:
    pages = {int(value) for value in re.findall(r"\bpg(\d{3})_sec\d{3}\.html?\b", reference, flags=re.IGNORECASE)}
    if re.search(r"(?:^|/)index\.html?\b", reference, flags=re.IGNORECASE):
        pages.add(1)
    if not pages and re.fullmatch(r"[\s,;&\d]+", reference):
        pages.update(int(value) for value in re.findall(r"\d{1,3}", reference))
    return sorted(page for page in pages if 1 <= page <= 168)


def main() -> None:
    document = Document(REPORT)
    table = document.tables[0]
    general_suggestions = []
    for paragraph in document.paragraphs:
        text = re.sub(r"\s+", " ", paragraph.text).strip()
        if not text:
            continue
        lowered = text.lower()
        if "nafasi" in lowered and ("kujibia" in lowered or "kujibu" in lowered):
            general_suggestions.append({"text": text, "status": "hold"})
        elif any(term in lowered for term in ["kurasa zote", "menu kuu", "consistency", "vielelezo vyote"]):
            general_suggestions.append({"text": text, "status": "global"})
    output = []
    for row_number, row in enumerate(table.rows[1:], 2):
        cells = ["\n".join(p.text for p in cell.paragraphs).strip() for cell in row.cells]
        status, category, note = classify(cells)
        pages = printed_pages(" ".join((cells[0], cells[1], cells[3])))
        if not pages:
            pages = printed_pages(cells[1], allow_trailing=True)
        candidates = [page + 6 for page in pages if page + 6 <= 168]
        direct_pages = reference_pages(cells[2])
        mapped_pages = direct_pages or candidates
        implementation_status = "pending_manual_verification"
        if status == "resolved_obsolete":
            implementation_status = "verified_by_pdf_faithful_rebuild"
        elif status == "hold":
            implementation_status = "held_by_user_instruction"
        elif category == "speech_cleanup":
            implementation_status = "implemented_global_rule"
        elif category == "speech_pronunciation" and "namba" in (cells[1] + " " + cells[3]).lower():
            implementation_status = "implemented_number_normalization_needs_page_qa"
        elif category == "book_wide":
            implementation_status = "implemented_global_rule_needs_full_qa"
        elif category in {"assistive_technology", "figure_description"}:
            implementation_status = "partially_implemented_needs_page_qa"
        output.append(
            {
                "matrix_row": row_number,
                "area": cells[0],
                "issue": cells[1],
                "old_adt_reference": cells[2],
                "recommendation": cells[3],
                "status": status,
                "category": category,
                "printed_page_references": pages,
                "candidate_pdf_pages": candidates,
                "direct_reference_pages": direct_pages,
                "mapped_pdf_pages": mapped_pages,
                "mapping_note": note,
                "implementation_status": implementation_status,
            }
        )

    summary = {
        status: sum(1 for item in output if item["status"] == status)
        for status in ["resolved_obsolete", "hold", "global", "page_specific"]
    }
    categories = {}
    for item in output:
        categories[item["category"]] = categories.get(item["category"], 0) + 1
    result = {
        "source": str(REPORT),
        "baseline_pages": 168,
        "summary": summary,
        "category_summary": dict(sorted(categories.items())),
        "general_suggestions": general_suggestions,
        "rows": output,
    }
    target = ROOT / "content" / "validation-matrix-mapping.json"
    target.write_text(json.dumps(result, ensure_ascii=False, indent=2) + "\n", encoding="utf-8")
    print(json.dumps(summary, ensure_ascii=False))


if __name__ == "__main__":
    main()
