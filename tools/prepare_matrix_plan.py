"""Turn the 11 Aug 2026 validation matrix into an auditable V1 work plan."""

from __future__ import annotations

import json
import re
from pathlib import Path

from docx import Document


ROOT = Path(__file__).resolve().parents[1]
REPORT = Path(r"C:\Users\Admin\Downloads\RIPOTI YA ADT VALIDATION SAYANSI 11.08.2026 (6).docx")
LEGACY_REPORT = Path(r"C:\Users\Admin\Downloads\RIPOTI YA ADT VALIDATION SAYANSI 11.08.2026.docx")


def category(text: str) -> tuple[str, str]:
    low = text.lower()
    if re.search(r"nafasi.{0,30}kuji(?:bu|bia)|sehemu.{0,20}kuji(?:bu|bia)|uingizaji.{0,20}majibu|sehemu.{0,20}kujaza", low):
        return "answer_space", "pending"
    if re.search(r"chemsha ?bongo|bangua ?bongo|\bqz\d+", low):
        return "quiz", "pending"
    if re.search(r"quorum|scratch|paint|pr(?:ogram|oram)(?:u)? said(?:izi|zi)|kipima joto sauti|saa ya mtetemo", low):
        return "assistive_technology", "pending"
    if re.search(r"kielelezo|picha|mchoro|maelezo ya sauti", low):
        return "figure_accessibility", "pending"
    if re.search(r"jumuishi|hisi|gusa|chunguza|baini|eleza|tambua|papasa", low):
        return "inclusive_language", "pending"
    if re.search(r"sauti|matamshi|tamk|taja|tawaja|kisoma skrini|lugha mchanganyiko", low):
        return "audio_pronunciation", "pending"
    if re.search(r"pangili|ihamishiwe|imejirudia|yamerudiwa|ukurasa usio sahihi", low):
        return "layout_order", "pending"
    return "content_correction", "pending"


def references(value: str) -> list[str]:
    found = re.findall(r"(?:adt/)?(?:pg\d+_sec\d+|qz\d+)\.html?", value, re.I)
    normalized = []
    for item in found:
        item = re.sub(r"^adt/", "", item, flags=re.I)
        item = re.sub(r"\.htm$", ".html", item, flags=re.I)
        normalized.append(item)
    return sorted(set(normalized))


doc = Document(REPORT)
table = doc.tables[0]
legacy_table = Document(LEGACY_REPORT).tables[0] if LEGACY_REPORT.exists() else None
items = []
for number, row in enumerate(table.rows[1:], 1):
    cells = [" ".join(cell.text.split()) for cell in row.cells]
    joined = " ".join(cells)
    kind, status = category(joined)
    refs = references(cells[2])
    # The revised matrix uses printed page numbers in its reference column.
    # Preserve the validator's earlier explicit HTML mapping for the same row.
    if not refs and legacy_table is not None and number < len(legacy_table.rows):
        legacy_reference = " ".join(legacy_table.rows[number].cells[2].text.split())
        refs = references(legacy_reference)
    items.append({
        "matrix_item": number,
        "area": cells[0],
        "issue": cells[1],
        "reference": cells[2],
        "recommendation": cells[3],
        "category": kind,
        "status": status,
        "files": refs,
        "missing_files": [name for name in refs if not (ROOT / name).exists()],
    })

general = [" ".join(p.text.split()) for p in doc.paragraphs if p.text.strip()][:35]
summary: dict[str, int] = {}
for item in items:
    summary[item["category"]] = summary.get(item["category"], 0) + 1
result = {
    "source": str(REPORT),
    "baseline": "Version 1 / b9b6648 content",
    "total_items": len(items),
    "policy": {
        "answer_spaces": "Implement for genuine learner-response prompts",
        "quizzes": "Remove standalone Chemsha bongo/Bangua bongo pages; preserve textbook exercises and make them interactive",
        "content": "Preserve visible textbook content unless a matrix row explicitly corrects it",
    },
    "category_summary": summary,
    "general_suggestions": general,
    "items": items,
}
(ROOT / "content" / "validation-matrix-plan.json").write_text(
    json.dumps(result, ensure_ascii=False, indent=2) + "\n", encoding="utf-8"
)
print(json.dumps({"total": len(items), "categories": summary}, ensure_ascii=False))
