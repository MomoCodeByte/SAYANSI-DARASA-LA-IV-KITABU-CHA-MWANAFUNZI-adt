from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


SOURCE = Path(r"C:\Users\Jacqueline\Downloads\SAYANSI DRS LA 4-MAREKEBISHO -.docx")
OUTPUT = Path(r"C:\Users\Jacqueline\Documents\ChatGPT\ADT-BOOKS-CONVERSION\SAYANSI-RIPOTI-YA-MAREKEBISHO.docx")


def clean(text):
    return " ".join(text.replace("\xa0", " ").split())


def shade_cell(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), fill)
    tc_pr.append(shading)


source = Document(SOURCE)
entries = []
for table in source.tables:
    for row in table.rows[1:]:
        values = [clean(cell.text) for cell in row.cells[:4]]
        if any(values):
            entries.append(values)

pages = sorted({entry[2] for entry in entries if entry[2]})
completed_pages = {
    "pg002_sec001.html",
    "pg003_sec001.html",
    "pg004_sec001.html",
    "pg005_sec001.html",
    "pg007_sec001.html",
    "pg009_sec001.html",
    "pg012_sec001.html",
}

doc = Document()
section = doc.sections[0]
section.top_margin = Inches(0.7)
section.bottom_margin = Inches(0.7)
section.left_margin = Inches(0.75)
section.right_margin = Inches(0.75)

styles = doc.styles
styles["Normal"].font.name = "Arial"
styles["Normal"].font.size = Pt(10.5)
styles["Title"].font.name = "Arial"
styles["Title"].font.size = Pt(18)
styles["Title"].font.bold = True
styles["Heading 1"].font.name = "Arial"
styles["Heading 1"].font.size = Pt(14)
styles["Heading 1"].font.bold = True
styles["Heading 1"].font.color.rgb = RGBColor(20, 91, 104)
styles["Heading 2"].font.name = "Arial"
styles["Heading 2"].font.size = Pt(11.5)
styles["Heading 2"].font.bold = True
styles["Heading 2"].font.color.rgb = RGBColor(20, 91, 104)

title = doc.add_paragraph(style="Title")
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
title.add_run("RIPOTI YA MAREKEBISHO YA KITABU CHA SAYANSI DARASA LA NNE")

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = subtitle.add_run("Nakala ya kazi ya kufuatilia utekelezaji wa maboresho ya ADT")
run.italic = True
run.font.size = Pt(11)

doc.add_paragraph()
summary = doc.add_table(rows=4, cols=2)
summary.style = "Table Grid"
summary_data = [
    ("Jina la kazi", "Marekebisho ya Sayansi Darasa la Nne"),
    ("Chanzo", "SAYANSI DRS LA 4-MAREKEBISHO -.docx"),
    ("Jumla ya hoja zilizorekodiwa", str(len(entries))),
    ("Kurasa zilizotajwa", str(len(pages))),
]
for row, (label, value) in zip(summary.rows, summary_data):
    row.cells[0].text = label
    row.cells[1].text = value
    shade_cell(row.cells[0], "D9EEF2")
    row.cells[0].paragraphs[0].runs[0].bold = True

doc.add_heading("1. Utangulizi", level=1)
doc.add_paragraph(
    "Ripoti hii imeandaliwa kwa ajili ya kuorodhesha na kufuatilia marekebisho "
    "yanayohitajika katika kitabu cha Sayansi Darasa la Nne kilichobadilishwa kuwa ADT. "
    "Marekebisho yanahusu matamshi ya sauti, mpangilio wa maandishi, picha na lebo, "
    "majedwali, maelezo ya vielelezo, pamoja na ulinganifu wa maudhui na kitabu cha PDF."
)

doc.add_heading("2. Lengo la marekebisho", level=1)
for text in [
    "Kuhakikisha maandishi na vichwa vinaonekana kwa mpangilio unaofanana katika kurasa zote.",
    "Kuhakikisha sauti inatamka maneno, herufi, namba na alama kwa usahihi.",
    "Kurekebisha picha, lebo na maelezo ya vielelezo bila kutoa majibu kwa mwanafunzi.",
    "Kuboresha mtiririko wa usomaji wa majedwali na sehemu za majaribio.",
    "Kuhakikisha toleo la kawaida na toleo rahisi kusoma yanafanya kazi kwa usahihi.",
]:
    doc.add_paragraph(text, style="List Bullet")

doc.add_heading("3. Orodha ya marekebisho kwa kila ukurasa", level=1)
for index, (number, issue, page, correction) in enumerate(entries, start=1):
    page_label = page or "Ukurasa haujatajwa"
    heading = doc.add_heading(f"3.{index}  {page_label}", level=2)

    paragraph = doc.add_paragraph()
    label = paragraph.add_run("Mapungufu yaliyoonekana: ")
    label.bold = True
    paragraph.add_run(issue or "Hayajaelezwa.")

    paragraph = doc.add_paragraph()
    label = paragraph.add_run("Marekebisho yanayohitajika: ")
    label.bold = True
    paragraph.add_run(correction or "Yafanyiwe marekebisho kulingana na maelekezo yaliyotolewa.")

    paragraph = doc.add_paragraph()
    label = paragraph.add_run("Hali ya utekelezaji: ")
    label.bold = True
    status_text = "Imekamilika na kuhakikiwa" if page in completed_pages else "Inasubiri kufanyiwa"
    status = paragraph.add_run(status_text)
    status.italic = True
    status.font.color.rgb = RGBColor(21, 128, 61) if page in completed_pages else RGBColor(160, 90, 0)

doc.add_heading("4. Utaratibu wa uhakiki baada ya marekebisho", level=1)
for text in [
    "Kupitia kila ukurasa uliotajwa na kulinganisha na maelekezo ya ripoti hii.",
    "Kusikiliza sauti za kawaida na sauti za toleo rahisi kusoma.",
    "Kuhakiki picha, lebo, vichwa, majedwali na mpangilio wa maandishi kwa mwonekano.",
    "Kufanya ukaguzi wa kitabu kizima ili kuthibitisha kuwa hakuna faili za sauti au muda wa sauti unaokosekana.",
    "Kuandaa preview kamili na muhtasari wa kila marekebisho yaliyotekelezwa.",
]:
    doc.add_paragraph(text, style="List Number")

doc.add_heading("5. Hitimisho", level=1)
doc.add_paragraph(
    "Marekebisho yatafanyika kwa kufuata hoja zote zilizoorodheshwa katika ripoti hii. "
    "Kila hoja itawekewa hali ya utekelezaji baada ya kukamilishwa na kuhakikiwa. "
    "Faili la awali lililotolewa halijabadilishwa; hii ni nakala mpya ya kazi inayoharirika."
)

for paragraph in doc.paragraphs:
    paragraph.paragraph_format.space_after = Pt(5)
    paragraph.paragraph_format.line_spacing = 1.08

doc.save(OUTPUT)
print(OUTPUT)
