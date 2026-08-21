from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt


SOURCE = Path(r"C:\Users\Jacqueline\Downloads\SAYANSI DRS LA 4-MAREKEBISHO -.docx")
OUTPUT = Path(r"C:\Users\Jacqueline\Documents\ChatGPT\ADT-BOOKS-CONVERSION\SAYANSI-MAREKEBISHO-NAKALA-YA-KAZI.docx")
COMPLETED_PAGES = {
    "pg002_sec001.html",
    "pg003_sec001.html",
    "pg004_sec001.html",
    "pg005_sec001.html",
    "pg007_sec001.html",
    "pg009_sec001.html",
    "pg012_sec001.html",
}


def copy_cell_text(source_cell, target_cell):
    target_cell.text = ""
    for index, source_paragraph in enumerate(source_cell.paragraphs):
        paragraph = target_cell.paragraphs[0] if index == 0 else target_cell.add_paragraph()
        paragraph.alignment = source_paragraph.alignment
        for source_run in source_paragraph.runs:
            run = paragraph.add_run(source_run.text)
            run.bold = source_run.bold
            run.italic = source_run.italic
            run.underline = source_run.underline


source = Document(SOURCE)
output = Document()

section = output.sections[0]
section.top_margin = Inches(0.55)
section.bottom_margin = Inches(0.55)
section.left_margin = Inches(0.55)
section.right_margin = Inches(0.55)

title = output.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
title_run = title.add_run("SAYANSI DARASA LA 4 – NAKALA YA KAZI YA MAREKEBISHO")
title_run.bold = True
title_run.font.size = Pt(14)

note = output.add_paragraph(
    "Hii ni nakala inayoharirika ya kufuatilia utekelezaji wa marekebisho. "
    "Faili la awali halijabadilishwa."
)
note.alignment = WD_ALIGN_PARAGRAPH.CENTER

for source_table in source.tables:
    table = output.add_table(rows=1, cols=5)
    table.style = "Table Grid"
    headers = ["Na", "Mapungufu", "Ukurasa", "Marekebisho", "Hali ya utekelezaji"]
    for cell, text in zip(table.rows[0].cells, headers):
        cell.text = text
        for run in cell.paragraphs[0].runs:
            run.bold = True

    for source_row in source_table.rows[1:]:
        target_cells = table.add_row().cells
        for source_cell, target_cell in zip(source_row.cells[:4], target_cells[:4]):
            copy_cell_text(source_cell, target_cell)
        page_name = source_row.cells[2].text.strip()
        target_cells[4].text = (
            "Imekamilika na kuhakikiwa"
            if page_name in COMPLETED_PAGES
            else "Inasubiri kufanyiwa"
        )

    widths = [0.45, 2.65, 0.75, 3.25, 1.15]
    for row in table.rows:
        for cell, width in zip(row.cells, widths):
            cell.width = Inches(width)

for paragraph in output.paragraphs:
    paragraph.paragraph_format.space_after = Pt(4)

for table in output.tables:
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.space_after = Pt(2)
                for run in paragraph.runs:
                    run.font.size = Pt(9)

output.save(OUTPUT)
print(OUTPUT)
